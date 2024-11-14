import os
import pandas as pd
from skimpy import clean_columns
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

folder_path = r'C:\Users\AkimFitzgerald\ORR Bulk Automation V2\Demo Folder that holds workbooks'

def sub_table(input_df, top_anchor, bottom_anchor):
    """
    Extracts a subset of rows from the input DataFrame based on the specified
    top and bottom anchors, and returns a DataFrame containing only the
    specified top anchor column with a concatenated string of the values
    in that column.

    Args:
        input_df (pandas.DataFrame): The input DataFrame.
        top_anchor (int or str): The value in the 'UC Portal A#' column that
            serves as the top anchor for the sub-table.
        bottom_anchor (int or str): The value in the 'UC Portal A#' column that
            serves as the bottom anchor for the sub-table.

    Returns:
        A pandas DataFrame containing only the specified top anchor column,
        with a concatenated string of the values in that column.

    Raises:
        ValueError: If the top or bottom anchor value is not found in the
            'UC Portal A#' column of the input DataFrame.
    """
    try:
        # Get the indices of the top and bottom anchor rows.
        top_idx = input_df.loc[input_df['UC Portal A#'] == top_anchor].index[0]
        bottom_idx = input_df.loc[input_df['UC Portal A#'] == bottom_anchor].index[0]
    except IndexError:
        # If the top or bottom anchor value is not found, raise a ValueError.
        raise ValueError("Top or bottom anchor value not found in 'UC Portal A#' column.")

    # Extract the sub-table and rename the columns.
    input_df = input_df.loc[top_idx + 1:bottom_idx - 1, :'UC Portal A#']
    input_df.columns = [top_anchor]

    # Concatenate the values in the top anchor column and return the result.
    input_df[top_anchor] = input_df[top_anchor].str.cat(sep=', ')
    input_df = input_df.head(1)
    return input_df

def extract_variables(input_df):
    """
    Given an input DataFrame, extracts the 'Employer:' and 'ID:' sub-tables
    and concatenates them into a new DataFrame. The 'sponsor_name' variable
    is also extracted from the sheet name.

    Args:
        input_df (pandas.DataFrame): The input DataFrame.

    Returns:
        A pandas DataFrame containing the 'Employer:' and 'ID:' sub-tables,
        concatenated into a new DataFrame.

    Raises:
        ValueError: If the sheet name cannot be parsed to extract the sponsor
            name, or if either of the sub-tables cannot be extracted.
    """

    try:
        # Extract the 'Employer:' and 'Identification Information:' sub-tables.
        employer_tab = sub_table(input_df, 'Employer:', 'Associates & Relatives:')
        id_tab = sub_table(input_df, 'Identification Information:', 'Employer:')
    except ValueError as e:
        # If either of the sub-tables cannot be extracted, raise a ValueError with a more specific message.
        raise ValueError(f"Unable to extract sub-tables from input DataFrame: {str(e)}")

    # Concatenate the sub-tables and return the result.
    df_concat = pd.concat([employer_tab, id_tab], axis=1)
    return df_concat

def extract_identifier_variables(input_df, row_index):
    """
    Given an input DataFrame and a row index, extracts several identifier
    variables and returns them as a tuple.

    Args:
        input_df (pandas.DataFrame): The input DataFrame.
        row_index (int): The row index from which to extract the variables.

    Returns:
        A tuple containing the following variables:
            - address (str): The address of the sponsor.
            - email (str): The email address of the sponsor.
            - phone (str): The phone number of the sponsor.
            - education (str): The school name of the sponsor.
            - uc_name (str): The name of the UC sponsor.
            - spon_name (str): The name of the sponsor.
            - ID (str): The UC Portal A# of the sponsor.

    Raises:
        ValueError: If any of the required variables cannot be extracted
            from the input DataFrame.
    """
    try:
        # Extract the address, email, phone, and education variables.
        address = str(input_df.iloc[row_index]['Last known address']) + ', ' + str(input_df.iloc[row_index]['City ']) + ', ' + str(input_df.iloc[row_index]['State'])
        email = str(input_df.iloc[row_index]['Sponsor Email Address'])
        phone = str(input_df.iloc[row_index]['Sponsor phone Number'])
        education = str(input_df.iloc[row_index]['UAM School Name'])
    except KeyError as e:
        # If any of the variables cannot be extracted, raise a ValueError with a more specific message.
        raise ValueError(f"Unable to extract required variables from input DataFrame: {str(e)}")

    try:
        # Extract the uc_name, spon_name, and ID variables.
        uc_name = f"{input_df.iloc[row_index]['UC Portal A#']}, {input_df.iloc[row_index]['First Name']} {input_df.iloc[row_index]['Last Name']}"
        spon_name = str(input_df.iloc[row_index]['Sponsor First Name']) + ' ' + str(input_df.iloc[row_index]['Sponsor Last Name'])
        ID = str(input_df.iloc[row_index]['UC Portal A#'])
    except KeyError as e:
        # If any of the variables cannot be extracted, raise a ValueError with a more specific message.
        raise ValueError(f"Unable to extract required variables from input DataFrame: {str(e)}")

    # Return the extracted variables as a tuple.
    return address, email, phone, education, uc_name, spon_name, ID

def extract_additional_identifier_variables(sheet_name, sheet_to_df_map):
    """
    Given a sheet name and a mapping of sheet names to DataFrames, extracts
    additional identifier variables from the specified sheet and returns
    them as a tuple.

    Args:
        sheet_name (str): The name of the sheet to extract variables from.
        sheet_to_df_map (dict): A dictionary mapping sheet names to DataFrames.

    Returns:
        A tuple containing the following variables:
            - employer (str): The employer of the sponsor.
            - id_info (str): Additional ID information for the sponsor.

    Raises:
        ValueError: If any of the required variables cannot be extracted
            from the specified sheet.
    """
    try:
        # Extract the DataFrame for the specified sheet.
        input_df = sheet_to_df_map[sheet_name]

        # Extract the employer and ID information from the DataFrame.
        df = extract_variables(input_df, sheet_name)
        employer = df.loc[df.index[0], 'Employer:']
        id_info = df.loc[df.index[0], 'Identification Information:']

    except KeyError as e:
        # If the specified sheet or any required variables cannot be extracted, raise a ValueError with a more specific message.
        raise ValueError(f"Unable to extract required variables from sheet '{sheet_name}': {str(e)}")

    # Return the extracted variables as a tuple.
    return employer, id_info

def generate_word_doc(input_df, folder_path, row_index=0):
    """
    Generate a Word document containing information from the input DataFrame.

    Args:
    input_df: A Pandas DataFrame containing information about a person.
    folder_path: The path of the folder where the generated Word document should be saved.
    row_index: The index of the row in the DataFrame containing the person's information to be used.

    Returns:
    None
    """
    
    # Get data from the sheet
    xls_1 = input_df
    # Extract identifier variables
    address = str(xls_1.iloc[row_index]['Last known address']) + ", " + str(xls_1.iloc[row_index]['City ']) + ", " + str(xls_1.iloc[row_index]['State'])
    email = xls_1.iloc[row_index]['Sponsor Email Address']
    phone = xls_1.iloc[row_index]['Sponsor Phone Number']
    education = xls_1.iloc[row_index]['UAM School Name ']
    uc_name = str(xls_1.iloc[row_index]['UC Portal A#']) + ", " + str(xls_1.iloc[row_index]['First Name']) + " " + str(xls_1.iloc[row_index]['Last Name'])
    uc_name_2 = xls_1.iloc[row_index]['First Name'].strip() + " " + xls_1.iloc[row_index]['Last Name'].strip()
    spon_name = str(xls_1.iloc[row_index]['Sponsor First Name']) + " " + str(xls_1.iloc[row_index]['Sponsor Last Name'])
    ID = xls_1.iloc[row_index]['UC Portal A#']
    # Get additional identifier variables from sheets
    identifier_2_df = extract_variables(input_df)
    employer = identifier_2_df.iloc[0]['Employer:']
    id_info = str(identifier_2_df.iloc[0]['Identification Information:']) if not pd.isna(identifier_2_df.iloc[0]['Identification Information:']) else ' '
    # Create new Word document
    document = Document()
    # Add header
    header = document.sections[0].header
    htable = header.add_table(1, 2, Inches(6))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture(r'C:\Users\AkimFitzgerald\Pictures\E24_Logo.png', width=Inches(2))
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add title and sponsor/UC information
    p_0 = document.add_paragraph('')
    run = p_0.add_run()
    run.add_break()
    run.add_break()
    p_0.add_run('Person’s Identification Project Summary Report').bold = True
    p_0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_1 = document.add_paragraph('')
    p_1.add_run('Sponsor: ').bold = True
    p_1.add_run(f'{spon_name}\t\t\t')
    p_1.add_run('UC: ').bold = True
    p_1.add_run(f'{uc_name}').bold = False
    run = p_1.add_run()
    run.add_break()
    p_2 = document.add_paragraph('')
    # Add tables
    p_2.add_run('TPG Provided Information:').bold = True
    p_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add UC table
    uc_table = document.add_table(2, 4)
    for j in range(1, 5):
        uc_table.cell(0, j-1).text = input_df.columns[j]
    for j in range(1, 5):
        uc_table.cell(1, j-1).text = str(input_df.values[row_index, j])
    p_x = document.add_paragraph('')
    # Add sponsor table
    sponsor_table = document.add_table(2, 4)
    for j in range(5, 9):
        sponsor_table.cell(0, j-5).text = input_df.columns[j]
    for j in range(5, 9):
        sponsor_table.cell(1, j-5).text = str(input_df.values[row_index, j])
    # Style tables
    uc_table.style = 'Light Grid Accent 1'
    sponsor_table.style = 'Light Grid Accent 1'
    # Add breaks
    p_3 = document.add_paragraph('')
    run = p_3.add_run()
    run.add_break()
    run.add_break()
    run.add_break()
    # Add research findings
    p_4 = document.add_paragraph('')
    p_4.add_run('Evolve24 Research Findings:').bold = True
    p_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_3.add_run()
    run.add_break()
    # Add address
    p_address = document.add_paragraph('')
    p_address.add_run(f'Address: ').bold = True
    p_address.add_run(f'{address}').bold = False
    # Add social media
    p_social_media = document.add_paragraph('')
    p_social_media.add_run('Social_Media:').bold = True
    # Add email
    p_email = document.add_paragraph('')
    p_email.add_run(f'Email: ').bold = True
    p_email.add_run(f'{email}').bold = False
    # Add employer
    p_employer = document.add_paragraph('')
    p_employer.add_run('Employer:').bold = True
    p_employer.add_run(f'{employer}').bold = False
    # Add phone number
    p_phone_number = document.add_paragraph('')
    p_phone_number.add_run(f'Phone Number: ').bold = True
    p_phone_number.add_run(f'{phone}').bold = False
    # Add education
    p_education = document.add_paragraph('')
    p_education.add_run(f'Education: ').bold = True
    p_education.add_run(f'{education}').bold = False
    # Add associates & relatives
    p_associates_relatives = document.add_paragraph('')
    p_associates_relatives.add_run('Associates & Relatives:').bold = True
    # Add identification information
    p_id_info = document.add_paragraph('')
    p_id_info.add_run('Identification Information:').bold = True
    p_id_info.add_run(f'{id_info}').bold = False
    # Add notes
    p_5 = document.add_paragraph('')
    p_5.add_run('Notes:').bold = True
    p_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_5.add_run()
    run.add_break()
    run.add_break()
    run.add_break()
    # Add footer
    footer_section = document.sections[0]
    footer = footer_section.footer
    footer_text = footer.paragraphs[0]
    footer_text.text = "Company Confidential & Company Proprietary | ©2022 evolve24"
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Save document
    document.save(f'{folder_path}/{ID}_{spon_name}.docx')

def generate_word_docs_in_folder(folder_path):
    """
    Given a folder path, iterates over each Excel workbook in the folder,
    loads it into a pandas DataFrame, and passes the DataFrame to the
    generate_word_doc function to output a Word doc.

    Args:
        folder_path (str): The path to the folder containing Excel workbooks.
        index (int): The index to pass to the generate_word_doc function.

    Returns:
        None.
    """

    # Iterate over each file in the folder.
    for filename in os.listdir(folder_path):
        # Check if the file is an Excel workbook.
        if filename.endswith('.xlsx') or filename.endswith('.xls'):
            # Load the Excel workbook into a pandas DataFrame.
            df = pd.read_excel(os.path.join(folder_path, filename), dtype={"Sponsor Phone Number": str})
            # Call the generate_word_doc function with the DataFrame.
            generate_word_doc(df, folder_path)

folder_path = r'C:\Users\AkimFitzgerald\ORR Bulk Automation V2\Demo Folder that holds workbooks'
generate_word_docs_in_folder(folder_path)